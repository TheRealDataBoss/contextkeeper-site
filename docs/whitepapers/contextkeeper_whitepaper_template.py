"""
ContextKeeper Whitepaper Template
==================================
Python equivalent of the Node.js docx build system used for WP-001 through
EXP-001. Produces the identical format: NAVY/ACCENT palette, Arial font,
US Letter, header/footer with page numbers, cover page, revision table,
body text, callout boxes, and data tables.

Usage:
    python contextkeeper_whitepaper_template.py

Output:
    ContextKeeper_Whitepaper_Template.docx

Dependencies:
    pip install python-docx

Customise the DOCUMENT_META dict and the build_content() function at the
bottom of this file for your specific paper.
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── PALETTE ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # Headings H1
ACCENT = RGBColor(0x2E, 0x75, 0xB6)   # Headings H2, rules, borders
DGRAY  = RGBColor(0x2C, 0x2C, 0x2C)   # Body text
MGRAY  = RGBColor(0x55, 0x55, 0x55)   # Cover sub-text, footer
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CALLOUT_BG = RGBColor(0xE8, 0xF1, 0xFB)   # Light blue callout background

# ── DOCUMENT META ─────────────────────────────────────────────────────────────
DOCUMENT_META = {
    "series_name":    "Enterprise Engineering White Paper Series",
    "title":          "Your Whitepaper Title Here",
    "subtitle":       "Optional Subtitle Line",
    "doc_id":         "WP-XXX",
    "version":        "1.0",
    "date":           "March 2026",
    "author":         "Steven Wazlavek",
    "org":            "ContextKeeper",
    "email":          "masterboss@contextkeeper.org",
    "website":        "contextkeeper.org",
    "classification": "Internal Engineering Reference",
    "header_short":   "ContextKeeper  |  Your Paper Title  |  WP-XXX v1.0",
}

# ── LOW-LEVEL XML HELPERS ─────────────────────────────────────────────────────

def set_run_font(run, size_pt: float, color: RGBColor | None = None,
                 bold: bool = False, italic: bool = False,
                 font_name: str = "Arial"):
    run.font.name        = font_name
    run.font.size        = Pt(size_pt)
    run.font.bold        = bold
    run.font.italic      = italic
    if color:
        run.font.color.rgb = color


def set_para_spacing(para, before_pt: float = 0, after_pt: float = 6,
                     line_spacing_pt: float | None = None):
    pf = para.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after  = Pt(after_pt)
    if line_spacing_pt:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line_spacing_pt)


def set_cell_shading(cell, fill_hex: str):
    """Fill a table cell background with a solid colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tcPr.append(shd)


def set_cell_borders(cell, color_hex: str = "CCCCCC", size: int = 4):
    """Apply single-line borders to all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex.upper())
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para_border_bottom(para, color_hex: str = "2E75B6", size: int = 6):
    """Add a bottom border to a paragraph (used for section rules)."""
    pPr   = para._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.upper())
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_para_border_top(para, color_hex: str = "2E75B6", size: int = 4):
    """Add a top border to a paragraph (used in footers)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    str(size))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color_hex.upper())
    pBdr.append(top)
    pPr.append(pBdr)


def add_page_number_field(para):
    """Insert a PAGE field run into a paragraph for auto page numbering."""
    run   = para.add_run()
    fldCh = OxmlElement("w:fldChar")
    fldCh.set(qn("w:fldCharType"), "begin")
    run._r.append(fldCh)

    run2   = para.add_run()
    instr  = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run2._r.append(instr)

    run3  = para.add_run()
    fldE  = OxmlElement("w:fldChar")
    fldE.set(qn("w:fldCharType"), "end")
    run3._r.append(fldE)

    for r in (run, run2, run3):
        set_run_font(r, 9, MGRAY)


def set_col_widths(table, widths_inches: list[float]):
    """Set exact column widths on a table."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_inches):
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement("w:tcW")
                tcW.set(qn("w:w"),    str(int(widths_inches[j] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


# ── PAGE SETUP ────────────────────────────────────────────────────────────────

def configure_page(doc: Document):
    """US Letter, 1-inch margins on all sides."""
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin   = Inches(1)
    section.bottom_margin = Inches(1)


# ── STYLES ────────────────────────────────────────────────────────────────────

def apply_styles(doc: Document):
    """
    Override Word's built-in Normal, Heading 1, and Heading 2 styles so all
    paragraphs inherit the correct font, colour, and spacing.
    """
    styles = doc.styles

    # Normal -- base for all body text
    normal = styles["Normal"]
    nf = normal.font
    nf.name  = "Arial"
    nf.size  = Pt(11)
    nf.color.rgb = DGRAY

    # Heading 1
    h1 = styles["Heading 1"]
    h1f = h1.font
    h1f.name  = "Arial"
    h1f.size  = Pt(18)
    h1f.bold  = True
    h1f.color.rgb = NAVY
    h1pf = h1.paragraph_format
    h1pf.space_before    = Pt(18)
    h1pf.space_after     = Pt(6)
    h1pf.page_break_before = True

    # Heading 2
    h2 = styles["Heading 2"]
    h2f = h2.font
    h2f.name  = "Arial"
    h2f.size  = Pt(14)
    h2f.bold  = True
    h2f.color.rgb = ACCENT
    h2pf = h2.paragraph_format
    h2pf.space_before = Pt(12)
    h2pf.space_after  = Pt(4)


# ── HEADER & FOOTER ───────────────────────────────────────────────────────────

def build_header(doc: Document, text: str):
    """Right-aligned italic header with ACCENT bottom border."""
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False

    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_para_border_bottom(para, color_hex="2E75B6", size=4)
    set_para_spacing(para, before_pt=0, after_pt=4)

    run = para.add_run(text)
    set_run_font(run, 9, MGRAY, italic=True)


def build_footer(doc: Document, suffix: str = "Internal Engineering Reference"):
    """Centred 'Page N | <suffix>' footer with ACCENT top border."""
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para_border_top(para, color_hex="2E75B6", size=4)
    set_para_spacing(para, before_pt=4, after_pt=0)

    r1 = para.add_run("Page ")
    set_run_font(r1, 9, MGRAY)

    add_page_number_field(para)

    r2 = para.add_run(f"  |  {suffix}")
    set_run_font(r2, 9, MGRAY)


# ── COVER PAGE ────────────────────────────────────────────────────────────────

def build_cover(doc: Document, meta: dict):
    """
    Centred cover page matching the ContextKeeper whitepaper style:
      - series name
      - ACCENT rule
      - title (large NAVY bold)
      - subtitle
      - ACCENT rule
      - doc ID / version / date
      - author block
      - classification
    """
    def cp(text: str, size: float, color: RGBColor, bold: bool = False,
           before_pt: float = 4, after_pt: float = 4) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=before_pt, after_pt=after_pt)
        run = p.add_run(text)
        set_run_font(run, size, color, bold=bold)

    def rule():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before_pt=0, after_pt=6)
        add_para_border_bottom(p, color_hex="2E75B6", size=6)

    # Vertical whitespace before content
    for _ in range(6):
        spacer = doc.add_paragraph()
        set_para_spacing(spacer, before_pt=0, after_pt=0)

    cp(meta["series_name"].upper(), 22, ACCENT, bold=True, before_pt=4, after_pt=4)
    rule()

    cp(meta["title"], 26, NAVY, bold=True, before_pt=8, after_pt=4)
    if meta.get("subtitle"):
        cp(meta["subtitle"], 18, NAVY, bold=True, before_pt=2, after_pt=4)

    rule()

    cp(f"{meta['doc_id']}  |  Version {meta['version']}  |  {meta['date']}",
       16, MGRAY, before_pt=8, after_pt=8)

    for _ in range(3):
        doc.add_paragraph()

    cp(f"Author: {meta['author']}",    16, MGRAY)
    cp(f"Organization: {meta['org']}", 16, MGRAY)
    cp(f"Contact: {meta['email']}",    16, MGRAY)
    cp(f"Web: {meta['website']}",      16, MGRAY)

    for _ in range(3):
        doc.add_paragraph()

    cp(f"Classification: {meta['classification']}", 16, MGRAY)

    doc.add_page_break()


# ── FRONT MATTER (copyright, revision table) ──────────────────────────────────

def fm_heading(doc: Document, text: str):
    """Bold small heading used in front matter (copyright, revision history)."""
    p = doc.add_paragraph()
    set_para_spacing(p, before_pt=12, after_pt=4)
    run = p.add_run(text)
    set_run_font(run, 12, NAVY, bold=True)
    return p


def build_front_matter(doc: Document, copyright_text: str,
                       authorship_text: str, revision_rows: list[list[str]]):
    """
    Adds the standard front-matter block:
      - Copyright notice
      - Statement of original authorship
      - Revision history table
    Then a page break.
    """
    fm_heading(doc, "COPYRIGHT AND INTELLECTUAL PROPERTY NOTICE")
    add_body_paragraph(doc, copyright_text)

    fm_heading(doc, "STATEMENT OF ORIGINAL AUTHORSHIP")
    add_body_paragraph(doc, authorship_text)

    fm_heading(doc, "REVISION HISTORY")
    add_table(
        doc,
        col_widths=[1.08, 1.42, 6.0],
        headers=["Version", "Date", "Changes"],
        rows=revision_rows,
    )
    doc.add_page_break()


# ── CONTENT HELPERS ───────────────────────────────────────────────────────────

def add_heading1(doc: Document, text: str, page_break: bool = True) -> None:
    """Section heading (H1). Adds a page break before by default."""
    p = doc.add_heading(text, level=1)
    p.paragraph_format.page_break_before = page_break


def add_heading2(doc: Document, text: str) -> None:
    """Sub-section heading (H2)."""
    doc.add_heading(text, level=2)


def add_body_paragraph(doc: Document, text: str,
                       before_pt: float = 0, after_pt: float = 6) -> None:
    """Standard body paragraph, 11pt Arial DGRAY."""
    p = doc.add_paragraph(style="Normal")
    set_para_spacing(p, before_pt=before_pt, after_pt=after_pt)
    run = p.add_run(text)
    set_run_font(run, 11, DGRAY)


def add_callout(doc: Document, text: str) -> None:
    """
    Shaded blue callout block matching the Node.js callout() helper.
    Uses a single-cell table for reliable shading.
    """
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8F1FB")
    set_cell_borders(cell, color_hex="2E75B6", size=4)

    # Set cell width to full content width (6.5 inches)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(6.5 * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    run = p.add_run(text)
    set_run_font(run, 11, NAVY, italic=True)

    # Spacer after callout
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before_pt=0, after_pt=4)


def add_table(doc: Document, col_widths: list[float],
              headers: list[str], rows: list[list[str]]) -> None:
    """
    Adds a styled data table.

    col_widths : list of column widths in inches (must sum to <= 6.5)
    headers    : list of header strings
    rows       : list of lists of cell strings
    """
    num_cols = len(col_widths)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style      = "Table Grid"
    table.alignment  = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_shading(cell, "1B2A4A")   # NAVY header
        set_cell_borders(cell, "2E75B6", 4)
        cell.paragraphs[0].clear()
        p    = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run  = p.add_run(hdr)
        set_run_font(run, 10, WHITE, bold=True)

    # Data rows
    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        fill = "F0F4FA" if i % 2 == 0 else "FFFFFF"  # alternating rows
        for j, text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            set_cell_shading(cell, fill)
            set_cell_borders(cell, "CCCCCC", 4)
            cell.paragraphs[0].clear()
            p    = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(4)
            run  = p.add_run(text)
            set_run_font(run, 10, DGRAY)

    # Apply column widths to every row
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < num_cols:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = OxmlElement("w:tcW")
                tcW.set(qn("w:w"),    str(int(col_widths[j] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

    # Spacer after table
    spacer = doc.add_paragraph()
    set_para_spacing(spacer, before_pt=0, after_pt=6)


def add_spacer(doc: Document, before_pt: float = 6) -> None:
    """Empty paragraph for vertical spacing."""
    p = doc.add_paragraph()
    set_para_spacing(p, before_pt=before_pt, after_pt=0)


# ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────────

def add_toc(doc: Document, entries: list[str]) -> None:
    """
    Adds a manual TOC page.
    entries : list of strings like "1. Executive Summary ..... 4"
    """
    add_heading1(doc, "Table of Contents", page_break=True)
    for entry in entries:
        p = doc.add_paragraph(style="Normal")
        set_para_spacing(p, before_pt=0, after_pt=3)
        run = p.add_run(entry)
        set_run_font(run, 11, DGRAY)
    doc.add_page_break()


# ── DOCUMENT ASSEMBLY ─────────────────────────────────────────────────────────

def build_document(meta: dict, content_fn) -> Document:
    """
    Full assembly pipeline:
      1. Create document
      2. Configure page (US Letter, 1-inch margins)
      3. Apply styles
      4. Build header + footer
      5. Build cover page
      6. Build front matter
      7. Call content_fn(doc) to add all sections
      8. Return the finished Document object
    """
    doc = Document()
    configure_page(doc)
    apply_styles(doc)
    build_header(doc, meta["header_short"])
    build_footer(doc, "Internal Engineering Reference")
    build_cover(doc, meta)

    build_front_matter(
        doc,
        copyright_text=(
            f"Copyright (c) 2026 {meta['author']}. All rights reserved. "
            "This document and the ideas, architectures, protocols, algorithms, "
            "and systems described herein constitute original intellectual property "
            f"of {meta['author']}, published under the {meta['org']} organization. "
            "No part of this document may be reproduced, distributed, transmitted, "
            "or used to create derivative works without the express written permission "
            "of the author, except for brief quotations in critical reviews and "
            "academic citations with proper attribution."
        ),
        authorship_text=(
            "The concepts, architectures, and systems described in this white paper "
            f"are the original work of {meta['author']}, developed during the design "
            f"and production engineering of the {meta['org']} platform "
            f"({meta['website']}). Third-party tools and research are cited explicitly."
        ),
        revision_rows=[
            [meta["version"], meta["date"], "Initial publication."],
        ],
    )

    content_fn(doc)
    return doc


# ── EXAMPLE CONTENT (replace with your own) ───────────────────────────────────

def build_content(doc: Document) -> None:
    """
    Replace everything in this function with your actual whitepaper content.
    This function demonstrates every available helper.
    """

    # ── TABLE OF CONTENTS ──────────────────────────────────────────────────
    add_toc(doc, [
        "1. Executive Summary ..... 4",
        "2. Background and Motivation ..... 5",
        "3. Architecture Overview ..... 6",
        "4. Data Tables Example ..... 7",
        "5. Conclusion ..... 8",
        "Appendix A. Glossary ..... 9",
        "References ..... 10",
    ])

    # ── SECTION 1 ─────────────────────────────────────────────────────────
    add_heading1(doc, "1. Executive Summary")
    add_body_paragraph(doc,
        "This document demonstrates the ContextKeeper whitepaper template. "
        "Replace this text with your executive summary. The template provides "
        "cover pages, front matter, H1/H2 headings, body paragraphs, callout "
        "boxes, and data tables -- all matching the standard ContextKeeper "
        "palette and formatting."
    )

    add_callout(doc,
        "Key finding or architectural principle goes here. Callout boxes are "
        "rendered with a light-blue background and NAVY italic text, matching "
        "the ContextKeeper highlight style."
    )

    add_body_paragraph(doc,
        "Continue with the narrative body text. Multiple paragraphs are added "
        "by calling add_body_paragraph() for each one. No bullet points -- "
        "write in dense technical paragraphs per the style guide."
    )

    # ── SECTION 2 ─────────────────────────────────────────────────────────
    add_heading1(doc, "2. Background and Motivation")
    add_body_paragraph(doc,
        "Section 2 body text. Describe the problem context, prior art, and "
        "motivation for the work in this section."
    )

    add_heading2(doc, "2.1 Sub-Section Example")
    add_body_paragraph(doc,
        "Sub-section text rendered as H2 in ACCENT blue. Use add_heading2() "
        "for all second-level headings."
    )

    # ── SECTION 3 ─────────────────────────────────────────────────────────
    add_heading1(doc, "3. Architecture Overview")
    add_body_paragraph(doc, "Architecture description.")

    add_heading2(doc, "3.1 Component Table")
    add_body_paragraph(doc,
        "The table below lists the primary components. Column widths are "
        "specified in inches and must sum to 6.5 or less."
    )
    add_spacer(doc, 6)
    add_table(
        doc,
        col_widths=[1.8, 2.4, 1.6, 0.7],
        headers=["Component", "Purpose", "Technology", "License"],
        rows=[
            ["Auth Layer",       "User management and session control",       "PHP 8 / MySQL",      "Proprietary"],
            ["Connector Bridge", "Uniform interface to external data sources", "Python 3.13",        "Proprietary"],
            ["Governance Layer", "DEL-v2 state machine enforcement",          "PHP / JSON-Schema",  "Proprietary"],
            ["Model Router",     "Multi-provider LLM dispatch",               "REST / OpenAI-compat","Proprietary"],
            ["Vector Store",     "Embedding storage for RAG retrieval",       "Qdrant / pgvector",  "Apache 2.0"],
        ],
    )

    # ── SECTION 4 ─────────────────────────────────────────────────────────
    add_heading1(doc, "4. Data Tables Example")
    add_body_paragraph(doc,
        "Tables support arbitrary column counts and widths. Header row is "
        "NAVY with white bold text. Data rows alternate light-blue and white."
    )
    add_spacer(doc, 6)
    add_table(
        doc,
        col_widths=[0.5, 2.6, 2.0, 1.4],
        headers=["#", "KPI Name", "SLO Target", "Escalation Trigger"],
        rows=[
            ["1",  "Probe Pass Rate (PPR)",        "100%",      "Any probe FAIL"],
            ["2",  "Initialization Pass Rate (IPR)", "100% (3/3)", "Any init FAIL"],
            ["3",  "Context Fabrication Rate (CFR)", "< 2%",     "> 10%: severe failure"],
            ["4",  "Schema Section Compliance",    "100%",      "Automatable check"],
            ["5",  "Handoff Continuity Rate (HCR)", "100% (3/3)", "State drop > 0"],
        ],
    )

    # ── SECTION 5 ─────────────────────────────────────────────────────────
    add_heading1(doc, "5. Conclusion")
    add_body_paragraph(doc,
        "Conclusion text. Summarise the contributions and next steps. "
        "Keep this section concise -- the substance is in the prior sections."
    )

    # ── APPENDIX ──────────────────────────────────────────────────────────
    add_heading1(doc, "Appendix A. Glossary")
    add_spacer(doc, 6)
    add_table(
        doc,
        col_widths=[2.0, 4.5],
        headers=["Term", "Definition"],
        rows=[
            ["DEL-v2",    "Deterministic Engineering Loop v2. ContextKeeper's formal state machine governance protocol."],
            ["SHAM",      "State, History, Action, Meaning. The four-field structured evidence schema used in EXP-001 run records."],
            ["Transport condition", "The mechanism (P/A/C) through which artifact files are made available to the model in a given run."],
        ],
    )

    # ── REFERENCES ────────────────────────────────────────────────────────
    add_heading1(doc, "References")
    refs = [
        "Wazlavek, S. (2026). DEL-v2 Technical Reference. ContextKeeper WP-001, v1.0. contextkeeper.org.",
        "Liu, N.F. et al. (2024). Lost in the Middle: How Language Models Use Long Contexts. TACL 12:157-173.",
        "Kohavi, R. et al. Trustworthy Online Controlled Experiments, Ch. 21: Sample Ratio Mismatch. Cambridge.",
    ]
    for ref in refs:
        add_body_paragraph(doc, ref, before_pt=0, after_pt=4)


# ── MAIN ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    output_path = "ContextKeeper_Whitepaper_Template.docx"
    doc = build_document(DOCUMENT_META, build_content)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables:     {len(doc.tables)}")
